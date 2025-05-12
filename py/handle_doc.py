from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from PIL import Image
from handle_image import CustomImage
from handle_log import log


def creat_docx(images: list[CustomImage], mode: int):
  """
  一頁兩張的函數
  :param images:
  :param mode:
  :return: 創建後的DOC檔
  """
  doc = Document()
  doc = set_font(doc)
  if mode == 1:
    for index, image in enumerate(images):
      doc = add_table_two_of_page_horizontal(doc, image, index + 1)
  elif mode == 2:
    for i in range(0, len(images), 2):
      doc = add_table_two_of_page_vertical(doc, images[i:i + 2], i + 1)
  elif mode == 6:
    for i in range(0, len(images), 3):
      doc = add_table_six_of_page(doc, images[i:i + 3], i + 1)
  return doc


def add_header(doc, title_text):
  """
  Word檔加入標頭
  :param doc: DOC檔
  :param title_text: 檔案標題文字
  :return: 編輯後的DOC檔
  """
  header = doc.sections[0].header
  title = header.paragraphs[0].add_run(title_text)
  title.font.size = Pt(20)
  title.font.bold = True
  header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
  log().debug('加入標頭成功')
  return doc


def set_font(doc):
  """
  用以設定預設字體及大小
  """
  try:
    doc.styles['Normal'].font.name = "Times New Roman"
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    log().debug('設定字體成功')
    return doc
  except Exception as e:
    log().exception(str(e))
    raise AttributeError('找不到指定字體')


def handle_number(no: int) -> str:
  """
  用以處理編號的函數
  :return: 處理完的文字
  """
  return f'0{no}' if no < 10 else str(no)


def add_table_two_of_page_horizontal(doc, image: CustomImage, index: int):
  """
  建立一頁2張圖片的表格(水平圖片，上下排佈）
  :param doc: 傳入的doc文件
  :param image: 傳入的圖片物件
  :param index: 序號，已經轉換成從1開始
  :return: 回傳doc文件
  """
  try:
    # 創建2*2的空表格
    table = doc.add_table(rows=2, cols=2, style='Table Grid')
    # 設定Row的高度
    table.rows[0].height = Cm(9)
    table.rows[1].height = Cm(2.2)
    # 設定Col的寬度
    for cell in table.columns[0].cells:
      cell.width = Cm(1.5)
    for cell in table.columns[1].cells:
      cell.width = Cm(15.5)

    # 合併表格
    table.cell(0, 0).merge(table.cell(0, 1))

    handle_table_write(
      table=table,
      image=image,
      index=index,
      image_cell=table.cell(0, 0),
      number_cell=table.cell(1, 0),
      remark_cell=table.cell(1, 1),
      max_height=9,
      max_width=15,
    )
    return doc

  except Exception as e:
    log().exception(str(e))


def add_table_two_of_page_vertical(doc, images: list[CustomImage], index: int):
  """
  建立一頁2張圖片的表格(垂直圖片，左右排佈）
  :param doc: 傳入的doc文件
  :param images: 傳入的圖片物件清單，最多會傳入兩張圖片
  :param index: 序號，已經轉換成從1開始
  :return: 回傳doc文件
  """
  try:
    # 創建2*2的空表格
    table = doc.add_table(rows=3, cols=2, style='Table Grid')
    # 設定Row的高度
    table.rows[0].height = Cm(18.2)
    table.rows[1].height = Cm(0.8)
    table.rows[2].height = Cm(3.5)
    # 設定Col的寬度
    for cell in table.columns[0].cells:
      cell.width = Cm(7.8)
    for cell in table.columns[1].cells:
      cell.width = Cm(7.8)

    handle_table_write(
      table=table,
      image=images[0],
      index=index,
      image_cell=table.cell(0, 0),
      number_cell=table.cell(1, 0),
      remark_cell=table.cell(2, 0),
      max_height=18,
      max_width=7.6,
    )
    if len(images) >= 2:
      handle_table_write(
        table=table,
        image=images[1],
        index=index + 1,
        image_cell=table.cell(0, 1),
        number_cell=table.cell(1, 1),
        remark_cell=table.cell(2, 1),
        max_height=18,
        max_width=7.6,
      )
    return doc

  except Exception as e:
    log().exception(str(e))


def add_table_six_of_page(doc, images: list[CustomImage], index):
  """
  建立一頁6張圖片的表格
  :param doc: 傳入的doc文件
  :param images: 傳入的圖片物件清單，最多會傳入三張圖片
  :param index: 序號，已經轉換成從1開始
  :return: 回傳doc文件
  """
  try:
    # 創建3*3的空表格
    table = doc.add_table(rows=3, cols=3, style='Table Grid')
    # 設定Row的高度
    table.rows[0].height = Cm(9)
    table.rows[1].height = Cm(0.8)
    table.rows[2].height = Cm(1.4)
    # 設定Col的寬度
    for i in range(0, 2):
      for cell in table.columns[i].cells:
        cell.width = Cm(5.2)

    handle_table_write(
      table=table,
      image=images[0],
      index=index,
      image_cell=table.cell(0, 0),
      number_cell=table.cell(1, 0),
      remark_cell=table.cell(2, 0),
      max_height=9,
      max_width=5,
    )
    if len(images) >= 2:
      handle_table_write(
        table=table,
        image=images[1],
        index=index + 1,
        image_cell=table.cell(0, 1),
        number_cell=table.cell(1, 1),
        remark_cell=table.cell(2, 1),
        max_height=9,
        max_width=5,
      )
    if len(images) >= 3:
      handle_table_write(
        table=table,
        image=images[2],
        index=index + 2,
        image_cell=table.cell(0, 2),
        number_cell=table.cell(1, 2),
        remark_cell=table.cell(2, 2),
        max_height=9,
        max_width=5,
      )

    return doc

  except Exception as e:
    log().exception(str(e))


def handle_table_write(table, image: CustomImage, index, image_cell, number_cell,
                       remark_cell, max_height: int | float, max_width: int | float):
  """
  寫入單一表格的資訊
  :param table: 表格
  :param image: 圖片物件
  :param index: 編號
  :param image_cell: 圖片的表格
  :param number_cell: 編號的表格
  :param remark_cell: 說明的表格
  :param max_height: 表格長上限
  :param max_width: 表格寬上限
  :return:
  """
  try:
    # 寫入文字及對齊
    number_cell.text = handle_number(index)  # 寫入編號
    number_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 表格文字垂直置中
    number_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 文字水平置中
    remark_cell.text = image.remark  # 寫入說明
    remark_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 文字水平置中
    remark_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP  # 表格文字垂直置頂

    # 處理圖片寬高
    img = Image.open(image.stream)
    default_aspect_ratio = max_width / max_height  # 預設寬高比
    image_aspect_ratio = img.width / img.height  # 圖片寬高比
    if image_aspect_ratio > default_aspect_ratio:
      # 圖片更扁，設定寬為表格上限
      image_cell.paragraphs[0].add_run().add_picture(image.stream, width=Cm(max_width))
    else:
      # 圖片更長，設定高為表格上限
      image_cell.paragraphs[0].add_run().add_picture(image.stream, height=Cm(max_height))

    # 設定圖片位置
    image_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    image_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    log().info(f'編號{index}圖片寫入完成')
  except Exception as e:
    log().error(f'編號{index}圖片寫入失敗，{str(e)}', exc_info=True)
