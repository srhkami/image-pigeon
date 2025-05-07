from docx import Document
from docx.shared import Inches, Cm
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from log import log
from image import CustomImage


def set_font(doc):
  '''
  用以設定預設字體及大小
  '''
  try:
    doc.styles['Normal'].font.name = "Times New Roman"
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    log().debug('設定預設字體成功')
  except:
    log().error('設定預設字體失敗', exc_info=True)
  return doc


def number(no):
  '''
  用以處理編號的函數
  :return: 處理完的文字
  '''
  if no < 10:
    return f'編號\n0{no}'
  else:
    return f'編號\n{no}'


def add_table(doc, image:CustomImage, index, show_type='H', mode=1):
  """
  :param doc: 傳入的doc文件
  :param image: 傳入的圖片物件
  :param index: 序號，已經轉換成從1開始
  :param show_type: 設定高或寬的模式
  :param mode: 如果為2，代表是多張模式
  :return: 回傳doc文件
  """
  # 創建表格
  try:
    # 創建5*3的空表格
    table = doc.add_table(rows=5, cols=3, style='Table Grid')

    # 設定Col的寬度
    for cell in table.columns[0].cells:
      cell.width = Cm(2)
    for cell in table.columns[1].cells:
      cell.width = Cm(12.5)
    for cell in table.columns[2].cells:
      cell.width = Cm(2.5)

    # 設定Row的高度
    table.rows[0].height = Cm(9)

    # 表格置中對齊
    # table.direction = WD_TABLE_ALIGNMENT.CENTER

    # 合併表格
    table.cell(0, 0).merge(table.cell(0, 2))
    table.cell(1, 0).merge(table.cell(4, 0))
    table.cell(1, 1).merge(table.cell(4, 2))
    # table.cell(4, 1).merge(table.cell(4, 2))
    # table.cell(3, 0).merge(table.cell(4, 0))
    # table.cell(3, 1).merge(table.cell(4, 1))

    # 寫入文字及對齊
    table.cell(1, 0).text = number(index)  ## 寫入編號
    table.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ## 表格文字垂直置中
    table.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  ## 文字水平置中
    # table.cell(1, 0).text = '攝影時間'
    # table.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # table.cell(1, 1).text = options['time']
    # table.cell(2, 0).text = '攝影地點'
    # table.cell(2, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # table.cell(2, 1).text = options['place']
    # table.cell(1, 1).text = '說明'
    # table.cell(3, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # table.cell(1, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.cell(1, 1).text = image.remark
    table.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    log().debug(f'編號{index} 表格創建成功')
  except:
    log().error(f'編號{index} 表格創建失敗', exc_info=True)

  # 寫入圖片
  try:
    if mode == 1:
      if show_type == 'H':
        table.cell(0, 0).paragraphs[0].add_run().add_picture(image.stream, height=Cm(9))
      else:
        table.cell(0, 0).paragraphs[0].add_run().add_picture(image.stream, width=Cm(15))
    # if mode == 2:
    #   # 如果mode是2，代表傳入的是清單，需循環取值
    #   for img in img_path:
    #     table.cell(0, 0).paragraphs[0].add_run().add_picture(img, height=Cm(9))
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    log().debug(f'編號{index} 圖片黏貼完成')
  except:
    log().error(f'編號{index} 圖片黏貼失敗', exc_info=True)

  return doc
