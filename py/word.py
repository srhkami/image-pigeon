from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image

from log import log
from doc import add_table, set_font


def one_of_one(title, images):
  '''
  一格單張的函數
  '''
  doc = Document()
  doc = set_font(doc)
  for index, image in enumerate(images):
    img = Image.open(image.stream)
    show_type = 'H'  # 預設設定為高9公分
    if img.width / img.height > 15 / 9:
      # 大於預設寬高比15：9，設定寬為15公分。
      show_type = 'W'
    doc = add_table(doc, image, index + 1, show_type=show_type)

  save(doc, title)


def save(doc, title_text):
  """

  :param doc: DOC檔
  :param title_text: 檔案標題文字
  :return:
  """
  header = doc.sections[0].header
  title = header.paragraphs[0].add_run(title_text)
  title.font.size = Pt(20)
  title.font.bold = True
  header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
  try:
    doc.save(rf'{title_text}.docx')
    log().info('執行完畢')
  except PermissionError:
    log().critical('有相同檔名的檔案未關閉', exc_info=True)
    raise PermissionError('儲存失敗，相同檔名的檔案未關閉')
  except Exception:
    log().critical('儲存失敗', exc_info=True)
    raise
